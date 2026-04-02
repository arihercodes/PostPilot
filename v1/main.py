import os
import base64
import datetime
import time
import json
from openai import OpenAI
import requests
from docx import Document
from docx.shared import Inches
from dotenv import load_dotenv

load_dotenv()

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

#brand data
def get_brand_data():
    print("\n" + "="*50)
    print("  BRAND SETUP")
    print("="*50)
    print("Let's set up your brand profile. Press Enter to skip any optional field.\n")

    brand_name = input("Brand Name: ").strip()
    tagline = input("Tagline: ").strip()
    brand_story = input("Brand Story (1-2 sentences): ").strip()
    founder_info = input("Founder Info (optional): ").strip()

    print("\nUnique Value Propositions (enter one per line, blank line to finish):")
    uvps = []
    while True:
        val = input("  UVP: ").strip()
        if not val:
            break
        uvps.append(val)

    print("\nBrand Values (enter one per line, blank line to finish):")
    values = []
    while True:
        val = input("  Value: ").strip()
        if not val:
            break
        values.append(val)

    print("\nBrand Accomplishments (enter one per line, blank line to finish):")
    accomplishments = []
    while True:
        val = input("  Accomplishment: ").strip()
        if not val:
            break
        accomplishments.append(val)

    print("\nProducts (enter one per line, blank line to finish):")
    products = []
    while True:
        val = input("  Product: ").strip()
        if not val:
            break
        products.append(val)

    print("\nServices (enter one per line, blank line to finish):")
    services = []
    while True:
        val = input("  Service: ").strip()
        if not val:
            break
        services.append(val)

    print("\nLogo Setup (optional but recommended):")
    print("  Provide the file path to your logo image (PNG or JPG).")
    print("  If skipped, the AI will generate a placeholder logo style instead.")
    while True:
        logo_path = input("  Logo file path (or press Enter to skip): ").strip()
        if not logo_path:
            logo_path = None
            break
        if os.path.exists(logo_path):
            break
        print(f"  [Error] File not found: '{logo_path}'. Please check the path and try again.")

    return {
        "brand_name": brand_name,
        "tagline": tagline,
        "brand_story": brand_story,
        "founder_info": founder_info,
        "Unique Value Proposition": uvps,
        "brand_values": values,
        "brand_accomplishments": accomplishments,
        "products": products,
        "services": services,
        "logo_path": logo_path,
    }


#config
def load_config():
    with open("config.json") as f:
        return json.load(f)


#brand goal
def get_user_goal():
    print("\n" + "="*50)
    print("  CONTENT GOAL")
    print("="*50)
    print("What do you want to generate today?")
    print("1. Branding")
    print("2. Product Promotion")
    print("3. Audience Engagement")
    print("4. Talent Acquisition")
    choice = input("Choose 1/2/3/4: ").strip()
    return {"1": "branding", "2": "product promotion", "3": "audience engagement", "4": "talent acquisition"}.get(choice, "branding")

#brand description
def get_brand_description(goal):
    print(f"\n" + "="*50)
    print(f"  POST DETAILS — {goal.upper()}")
    print("="*50)

    if goal == "branding":
        audience = input("1. Who is your primary audience? ")
        tone = input("2. What tone should your brand convey? (e.g., Funny, Professional, Empathetic etc.) ")
        mission = input("3. What is your mission or long-term vision? ")
        image_style = input("4. What type of post do you want to make? (Brand story post, founder intro post, Unique Value Post, Brand promise post, brand milestones post?) ")
        description = f"Audience: {audience} Tone: {tone}. Mission: {mission}. post_type: {image_style}."
        return description, image_style

    elif goal == "product promotion":
        product = input("1. What product are you promoting? ")
        benefits = input("2. What are the top benefits of this product? ")
        audience = input("3. Who is this product mainly for? ")
        usage = input("4. How is the product used in daily life? ")
        value = input("5. Why should customers buy your product? ")
        tone = input("6. What tone should your brand convey? (e.g., Funny, Professional, Empathetic etc.) ")
        image_style = input("7. What type of post do you want to make? (Product selling post, how-to-use post, unique selling point post etc.?) ")
        description = f"Product: {product}. Benefits: {benefits}. Audience: {audience}. Usage: {usage}. Value: {value}. Tone: {tone}. post_type: {image_style}."
        return description, image_style

    elif goal == "audience engagement":
        audience_type = input("1. What type of audience are you targeting? (e.g., rural/urban people, stressed professionals, fitness lovers, old people etc.) ")
        audience_age = input("2. What age group are you targeting? ")
        tone = input("3. What tone should your brand convey? (e.g., Funny, Professional, Empathetic etc.) ")
        image_style = input("4. What format should the post take? (e.g., poll(this/that), question post, meme post, product giveaway post, would you rather post) ")
        engagement_type = input("5. What kind of engagement do you want? (e.g., likes, comments, shares, dms, saves) ")
        description = f"Audience_Type: {audience_type}. Audience_Age: {audience_age}. Tone: {tone}. post_type: {image_style}. Engagement_Type: {engagement_type}"
        return description, image_style

    elif goal == "talent acquisition":
        position = input("1. What position are you hiring for? ")
        team = input("2. Which type of team/department is the position for? ")
        skills = input("3. Key skills or experience required? ")
        tone = input("4. What tone should this hiring post convey? (e.g., Inspiring, Friendly, Professional etc.) ")
        image_style = input("5. What type of post do you want? (e.g., Join Us post, We Are Hiring banner, Open Roles showcase) ")
        description = f"Hiring for: {position} in {team} team. Required: {skills}. Tone: {tone}. Post Type: {image_style}"
        return description, image_style

    else:
        return "Please choose a correct option.", ""


#post type info
def extract_post_specific_info(post_type, brand_data):
    pt = post_type.lower()
    if "unique selling" in pt or "unique value" in pt:
        return "Unique Selling Points: " + "; ".join(brand_data["Unique Value Proposition"])
    elif "brand story" in pt:
        return "Brand Story: " + brand_data["brand_story"]
    elif "founder" in pt:
        return "Founder Insight: " + brand_data["founder_info"]
    elif "milestone" in pt:
        return "Brand Milestones: " + "; ".join(brand_data["brand_accomplishments"])
    elif "brand value" in pt:
        return "Brand Values: " + "; ".join(brand_data["brand_values"])
    elif "product selling" in pt:
        return "Products we offer: " + "; ".join(brand_data["products"])
    elif "services" in pt:
        return "Services we offer: " + "; ".join(brand_data["services"])
    else:
        return "Key Info: " + brand_data["tagline"]

#news
def fetch_news_headline():
    api_key = os.getenv("GNEWS_API_KEY")
    url = f"https://gnews.io/api/v4/top-headlines?lang=en&country=in&max=1&apikey={api_key}"
    response = requests.get(url)
    if response.status_code == 200:
        articles = response.json().get("articles", [])
        if articles:
            return articles[0]['title']
    return "No news headline found"

#events
def fetch_upcoming_event():
    today = datetime.date.today().strftime("%B %d, %Y")
    prompt = (
        f"Today is {today}. Suggest one upcoming Indian event or holiday that is on or after today (not before). "
        f"It must occur within the next 14 days. Format the response like this: 'Event Name'."
    )
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}]
    )
    print(response.usage)
    return response.choices[0].message.content.strip()


#post ideas
def generate_post_ideas(relevant_info, user_goal, brand_description, brand_data, selected_topic, post_type):
    system_prompt = (
        f"You are a creative director for a wellness brand called '{brand_data['brand_name']}'.\n"
        f"The brand offers: {', '.join(brand_data['services'])} and products like: {', '.join(brand_data['products'])}.\n\n"
        f"Your job is to create **static Instagram post ideas** (not reels or stories).\n"
        f"Each idea must be suitable for a '{post_type}' format.\n"
        f"Here is more context based on post_type '{post_type}':\n{relevant_info}\n"
        f"The goal of this content is: '{user_goal}'.\n"
        f"The post must relate to this event/topic: {selected_topic}.\n"
        f"Use the brand description and tone given in {brand_description} to guide your suggestions:\n\n"
        f"ONLY return 3 short Instagram post **titles** (do not describe the image or explain anything).\n"
        f"Make sure each idea fits both the topic AND the post type.\n"
        f"Return the titles as a simple numbered list of strings.\n"
    )
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": system_prompt}]
    )
    print(response.usage)
    post_text = response.choices[0].message.content.strip()
    post_ideas = [line for line in post_text.split('\n') if line.strip().startswith(('1.', '2.', '3.'))]
    return post_ideas, system_prompt

#images 
def generate_images(relevant_info, brand_data, system_prompt, selected_idea, config, post_type, n=3):
    folder = "images_generated"
    os.makedirs(folder, exist_ok=True)
    image_paths = []
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")

    logo_path = brand_data.get("logo_path")
    logo_base64 = None
    logo_media_type = "image/png"
    if logo_path and os.path.exists(logo_path):
        ext = os.path.splitext(logo_path)[1].lower()
        logo_media_type = "image/jpeg" if ext in (".jpg", ".jpeg") else "image/png"
        with open(logo_path, "rb") as f:
            logo_base64 = base64.b64encode(f.read()).decode("utf-8")
        logo_instruction = f"- Place the brand logo (shown in the reference image) at the {config['logo_position']}. Reproduce it faithfully — do not alter its shape or design."
    else:
        logo_instruction = f"- The brand has no logo uploaded. Leave the {config['logo_position']} clean or add a simple minimal placeholder that fits the brand aesthetic."

    for i in range(n):
        filename = f"{brand_data['brand_name'].lower().replace(' ', '_')}_image_{timestamp}_{i+1}.png"
        image_path = os.path.join(folder, filename)

        prompt_text = (
            f"{system_prompt}\n\n"
            f"You are designing a static Instagram image for the brand '{brand_data['brand_name']}'.\n"
            f"Post Type: {post_type}\n"
            f"Post Idea Title: '{selected_idea}'\n"
            f"Base your image ONLY on this info:\n{relevant_info}\n\n"
            f"Use relevant colours to make the image eye-catching.\n"
            f"{logo_instruction}\n"
            f"- Add the brand name '{brand_data['brand_name']}' at {config['text_position']}.\n"
            f"- Style must match the theme: {post_type}.\n"
            f"- Maintain an aspect ratio of {config['aspect_ratio']}.\n"
            f"- Avoid listing generic services or all products unless that is the post type.\n"
        )

        if logo_base64:
            api_input = [
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "input_image",
                            "source": {
                                "type": "base64",
                                "media_type": logo_media_type,
                                "data": logo_base64,
                            },
                        },
                        {"type": "input_text", "text": prompt_text},
                    ],
                }
            ]
        else:
            api_input = prompt_text

        response = client.responses.create(
            model="gpt-4o",
            input=api_input,
            tools=[{"type": "image_generation"}]
        )
        print(response.usage)

        image_data = [output.result for output in response.output if output.type == "image_generation_call"]
        if image_data:
            image_base64 = image_data[0]
            with open(image_path, "wb") as f:
                f.write(base64.b64decode(image_base64))
            print(f"[Image Saved] {image_path}")
            image_paths.append(image_path)
        else:
            print(f"[Error] No image generated for prompt {i+1}.")

    return image_paths

#captions
def generate_captions(image_topic, image_paths, brand_name):
    captions = []
    for i, path in enumerate(image_paths):
        with open(path, "rb") as image_file:
            image_data = image_file.read()
        base64_image = base64.b64encode(image_data).decode("utf-8")

        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "image_url",
                            "image_url": {"url": f"data:image/png;base64,{base64_image}"}
                        },
                        {
                            "type": "text",
                            "text": (
                                f"This is an Instagram image for a brand called {brand_name}. "
                                f"Topic: {image_topic}. Generate a creative, relatable caption in max 20 words, "
                                f"suitable for this image. Add 3-4 popular and relevant hashtags."
                            )
                        }
                    ]
                }
            ]
        )
        print(response.usage)
        caption = response.choices[0].message.content.strip()
        print(f"\n[Caption {i+1}] {caption}")
        captions.append(caption)
    return captions


#generate document with all results
def generate_summary_docx(brand_name, topic, user_goal, post_type, ideas, selected_idea, captions, image_paths):
    doc = Document()
    doc.add_heading(f"{brand_name} — Instagram Content Summary", level=1)
    doc.add_paragraph(f"Topic: {topic}")
    doc.add_paragraph(f"Goal: {user_goal}")
    doc.add_paragraph(f"Post Type: {post_type}")

    doc.add_heading("Post Ideas", level=2)
    for idea in ideas:
        doc.add_paragraph(idea, style='List Number')

    doc.add_heading("Selected Idea", level=2)
    doc.add_paragraph(selected_idea)

    if captions:
        doc.add_heading("Generated Captions", level=2)
        for i, caption in enumerate(captions):
            doc.add_paragraph(f"Caption {i+1}: {caption}")

    doc.add_heading("Images", level=2)
    for path in image_paths:
        if os.path.exists(path):
            doc.add_picture(path, width=Inches(4))
            doc.add_paragraph(path)

    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{brand_name.lower().replace(' ', '_')}_post_summary_{timestamp}.docx"
    doc.save(filename)
    print(f"\n[Document Saved] {filename}")


#publishing
def upload_to_catbox(image_path):
    print("Uploading image to Catbox.moe...")
    try:
        with open(image_path, "rb") as f:
            response = requests.post(
                "https://catbox.moe/user/api.php",
                data={"reqtype": "fileupload"},
                files={"fileToUpload": f}
            )
            if response.status_code == 200:
                return response.text.strip()
            else:
                print("[Error] Upload failed", response.status_code)
                return None
    except Exception as e:
        print("[Exception]", e)
        return None

#instagram
def publish_to_instagram(image_url, caption, access_token):
    accounts = requests.get(
        'https://graph.facebook.com/v23.0/me/accounts',
        params={'access_token': access_token}
    ).json()
    page_id = accounts['data'][0]['id']

    ig_data = requests.get(
        f'https://graph.facebook.com/v23.0/{page_id}',
        params={'fields': 'instagram_business_account', 'access_token': access_token}
    ).json()
    ig_id = ig_data['instagram_business_account']['id']

    creation_res = requests.post(
        f"https://graph.facebook.com/v23.0/{ig_id}/media",
        data={"image_url": image_url, "caption": caption, "access_token": access_token}
    ).json()

    creation_id = creation_res.get("id")
    if creation_id:
        time.sleep(3)
        publish_res = requests.post(
            f"https://graph.facebook.com/v23.0/{ig_id}/media_publish",
            data={"creation_id": creation_id, "access_token": access_token}
        ).json()
        print("Instagram Post Response:", publish_res)
    else:
        print("[Error] Media container creation failed:", creation_res)

#facebook
def publish_to_facebook(image_url, caption, access_token):
    accounts = requests.get(
        'https://graph.facebook.com/v23.0/me/accounts',
        params={'access_token': access_token}
    ).json()
    page_id = accounts['data'][0]['id']

    response = requests.post(
        f"https://graph.facebook.com/v23.0/{page_id}/photos",
        data={'url': image_url, 'caption': caption, 'access_token': access_token}
    )
    print("Facebook Post Response:", response.json())

#main
def main():
    try:
        brand_data = get_brand_data()
        user_goal = get_user_goal()
        config = load_config()
        brand_description, post_type = get_brand_description(user_goal)

        event = fetch_upcoming_event()
        news = fetch_news_headline()

        print(f"\nChoose your topic:")
        print(f"1. Upcoming Event → {event}")
        print(f"2. News Headline  → {news}")
        print(f"3. None of these – Enter a custom topic")
        topic_choice = input("Select 1 / 2 / 3: ").strip()

        if topic_choice == "1":
            selected_topic = event
        elif topic_choice == "2":
            selected_topic = news
        else:
            selected_topic = input("Enter a custom topic (e.g., Stress Awareness, Sleep Rituals, etc.): ").strip()

        print(f"\nSelected Topic: {selected_topic}")

        relevant_info = extract_post_specific_info(post_type, brand_data)
        ideas, system_prompt = generate_post_ideas(relevant_info, user_goal, brand_description, brand_data, selected_topic, post_type)

        print("\nHere are 3 post ideas:")
        for idea in ideas:
            print(idea.strip())

        choice = input("\nWhich idea do you want to generate images for? (1/2/3): ").strip()
        if not choice.isdigit() or int(choice) not in [1, 2, 3]:
            raise ValueError("Invalid choice. Please enter 1, 2, or 3.")

        selected_idea = ideas[int(choice) - 1]
        image_paths = generate_images(relevant_info, brand_data, system_prompt, selected_idea, config, post_type)

        captions = []
        gen_caption = input("\nDo you want to generate captions for the images? (yes/no): ").strip().lower()
        if gen_caption == "yes":
            captions = generate_captions(selected_idea, image_paths, brand_data["brand_name"])

        generate_summary_docx(
            brand_name=brand_data["brand_name"],
            topic=selected_topic,
            user_goal=user_goal,
            post_type=post_type,
            ideas=ideas,
            selected_idea=selected_idea,
            captions=captions,
            image_paths=image_paths
        )

        print("\nGenerated Images:")
        for idx, path in enumerate(image_paths):
            print(f"{idx+1}. {path}")

        selected_img_idx = int(input("Select image to upload (1/2/3): ").strip()) - 1
        selected_img_path = image_paths[selected_img_idx]

        if captions:
            print("\nGenerated Captions:")
            for idx, caption in enumerate(captions):
                print(f"{idx+1}. {caption}")
            selected_caption_idx = int(input("Select caption to post (1/2/3): ").strip()) - 1
            selected_caption = captions[selected_caption_idx]
        else:
            selected_caption = input("Enter your custom caption: ").strip()

        image_url = upload_to_catbox(selected_img_path)
        print(f"\nPublic Image URL: {image_url}")
        if not image_url:
            print("[Error] Image upload failed. Exiting.")
            return

        platform = input("\nWhere do you want to post? (instagram/facebook): ").strip().lower()

        if platform == "instagram":
            token = os.getenv("INSTAGRAM_ACCESS_TOKEN")
            publish_to_instagram(image_url, selected_caption, token)
        elif platform == "facebook":
            token = os.getenv("FACEBOOK_ACCESS_TOKEN")
            publish_to_facebook(image_url, selected_caption, token)
        else:
            print("Invalid platform selected.")

    except Exception as e:
        print(f"\n[Error] {str(e)}")


if __name__ == "__main__":
    main()